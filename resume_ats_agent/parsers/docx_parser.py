import os
from resume_ats_agent.parsers.base import BaseResumeParser
from resume_ats_agent.core.exceptions import ParserError
from resume_ats_agent.core.logger import logger

class DocxParser(BaseResumeParser):
    """Concrete parser strategy to extract text content from DOCX resume files."""
    
    def parse(self, file_path: str) -> str:
        if not os.path.exists(file_path):
            raise ParserError(f"File not found: {file_path}")
            
        logger.info(f"Parsing DOCX file: {file_path}")
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text:
                    full_text.append(para.text)
            
            # Extract table cells content (common in resumes)
            for table in doc.tables:
                for row in table.rows:
                    row_content = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_content:
                        full_text.append(" | ".join(set(row_content)))  # simple CSV representation
                        
            text = "\n".join(full_text)
            if text.strip():
                logger.info("DOCX parsed successfully.")
                return text
        except Exception as e:
            raise ParserError(f"Failed to parse DOCX file '{file_path}': {str(e)}")
            
        raise ParserError(f"No textual content could be extracted from DOCX file: {file_path}")
